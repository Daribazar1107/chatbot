"""
ingest.py — NUM chatbot data ingestion (v3)

Reads all data files and builds Pinecone vector index + BM25 disk cache.

SUPPORTED FILES AND THEIR STRUCTURES:
  chuluu.json   — single object with "provisions" array  {group, tags, provisions:[{provision, content}]}
  courses.json  — array of course objects                [{Course_Index, Course_Name, Credit_Hours, Level, Brief_Content, Department}]
  grading.json  — array of typed objects                 [{id, type, title, content}]
                  types: table | definitions | regulation_section | standards | methodology
  level.json    — single regulation_table object         {type:"regulation_table", content:[{program, levels:[{name, credit_hours}]}]}
  schedule.json — array of schedule objects              [{instructor, course, index, semester, day, time, room, text, source}]
  teachers.json — array of teacher objects               [{name, department, room_number, phone, email}]
  tuition.json  — array of tuition objects               [{type, school, bank, account_number, text} | {academic_year, enrollment_period, school, tuition:{general_foundation, major_foundation}}]
"""

import os, re, json, csv, time, hashlib
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from pinecone import Pinecone, ServerlessSpec
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv

load_dotenv()

# ── CONFIGURATION ────────────────────────────────────────────────────────────
PINECONE_KEY  = os.getenv("PINECONE_API_KEY")
INDEX_NAME    = "muis-chatbot"   # must match INDEX_NAME in app.py
DATA_FOLDER   = "data"
CHUNK_SIZE    = 700              # max characters per chunk before splitting
OVERLAP_SIZE  = 120              # overlap between split chunks for context continuity
EMBED_MODEL   = "all-mpnet-base-v2"  # must match EMBED_MODEL in app.py
EMBED_DIM     = 768              # output dimension of all-mpnet-base-v2
BATCH_SIZE    = 50               # vectors per Pinecone upsert batch
MIN_CHUNK_LEN = 30               # discard chunks shorter than this

# ── CLIENTS ──────────────────────────────────────────────────────────────────
pc = Pinecone(api_key=PINECONE_KEY)
print(f"🔧 Loading embedder: {EMBED_MODEL}")
embedder = SentenceTransformer(EMBED_MODEL)
print("✅ Embedder ready.")


# ── PINECONE INDEX ───────────────────────────────────────────────────────────
def get_or_create_index():
    """Create the Pinecone index if it doesn't exist; recreate if dimension mismatch."""
    existing = [i.name for i in pc.list_indexes()]
    if INDEX_NAME in existing:
        desc = pc.describe_index(INDEX_NAME)
        current_dim = desc.dimension
        if current_dim != EMBED_DIM:
            print(f"⚠️  Dimension mismatch ({current_dim} ≠ {EMBED_DIM}). Recreating index...")
            pc.delete_index(INDEX_NAME)
            time.sleep(3)
        else:
            print(f"✅ Index '{INDEX_NAME}' ready (dim={current_dim}).")
            return pc.Index(INDEX_NAME)

    print(f"🆕 Creating index '{INDEX_NAME}' (dim={EMBED_DIM})...")
    pc.create_index(
        name=INDEX_NAME,
        dimension=EMBED_DIM,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1"),
    )
    time.sleep(5)
    print("✅ Index created.")
    return pc.Index(INDEX_NAME)


# ── CHUNK ID ─────────────────────────────────────────────────────────────────
def make_id(text: str, idx: int) -> str:
    """Stable unique ID: position + MD5 prefix of text content."""
    h = hashlib.md5(text.encode()).hexdigest()[:8]
    return f"doc-{idx:05d}-{h}"


# ── TEXT SPLITTER ─────────────────────────────────────────────────────────────
def split_chunks(text: str, source: str) -> list[dict]:
    """
    Split long text into overlapping chunks at natural boundaries.
    Tries to split at section numbers, sentence endings, or fixed size.
    """
    text = re.sub(r"\n{3,}", "\n\n", text.strip())
    section_breaks = re.compile(r'(?=\n\d+\.\d+[\.\d]*[\s\.])')
    sections = section_breaks.split(text)
    chunks = []

    for section in sections:
        section = section.strip()
        if not section:
            continue
        if len(section) <= CHUNK_SIZE:
            if len(section) >= MIN_CHUNK_LEN:
                chunks.append({"text": section, "source": source})
        else:
            start = 0
            while start < len(section):
                end   = start + CHUNK_SIZE
                chunk = section[start:end]
                if end < len(section):
                    for sep in [".\n", ".\r", ".\n\n", ". "]:
                        pos = chunk.rfind(sep)
                        if pos > CHUNK_SIZE // 2:
                            end   = start + pos + len(sep)
                            chunk = section[start:end]
                            break
                chunk = chunk.strip()
                if len(chunk) >= MIN_CHUNK_LEN:
                    chunks.append({"text": chunk, "source": source})
                next_start = end - OVERLAP_SIZE
                if next_start <= start:
                    break
                start = next_start

    return chunks


# ── PARSER: chuluu.json ───────────────────────────────────────────────────────
def parse_chuluu(item: dict, src: str) -> list[dict]:
    """
    chuluu.json structure:
      { group: str, tags: [str], provisions: [{provision: str, content: str}] }

    Each provision becomes one chunk, tagged with the group name and tags.
    """
    results = []
    group = item.get("section", item.get("group", "")).strip()  # e.g. "FIVE. Taking Leave..."
    tags  = item.get("tags", [])                                 # e.g. ["leave", "dismissal"]

    for prov in item.get("provisions", []):
        if not isinstance(prov, dict):
            continue
        clause  = prov.get("clause", prov.get("provision", "")).strip()  # e.g. "5.1.1"
        content = prov.get("content", "").strip()                         # the regulation text
        if not content:
            continue

        parts = []
        if group:   parts.append(f"Section: {group}.")
        if clause:  parts.append(f"Clause {clause}:")
        parts.append(content)
        if tags:    parts.append(f"Related: {', '.join(tags)}.")

        text = " ".join(parts)
        if len(text) >= MIN_CHUNK_LEN:
            results.append({"text": text, "source": src})

    return results


# ── PARSER: courses.json ──────────────────────────────────────────────────────
def parse_course(item: dict, src: str) -> list[dict]:
    """
    courses.json structure:
      { Course_Index, Course_Name, Credit_Hours, Level, Brief_Content, Department }
    """
    idx   = item.get("Course_Index", "").strip()
    name  = item.get("Course_Name",  "").strip()
    creds = item.get("Credit_Hours", "")
    level = item.get("Level",        "").strip()
    desc  = item.get("Brief_Content","").strip()
    dept  = item.get("Department",   "").strip()

    if not name:
        return []

    parts = [f"Course: {name}" + (f" ({idx})" if idx else "") + "."]
    if level:  parts.append(f"Level: {level}.")
    if creds:  parts.append(f"Credit hours: {creds}.")
    if dept:   parts.append(f"Department: {dept}.")
    if desc:   parts.append(f"Description: {desc}")

    text = " ".join(parts)
    if len(text) > CHUNK_SIZE:
        return split_chunks(text, src)
    if len(text) >= MIN_CHUNK_LEN:
        return [{"text": text, "source": src}]
    return []


# ── PARSER: grading.json ──────────────────────────────────────────────────────
def parse_grading(item: dict, src: str) -> list[dict]:
    """
    grading.json — array of typed objects. Each type is parsed differently:

      type="table"              → score/letter/grade_point rows
      type="definitions" list   → special grade notations (W, WF, I, R, F, NR, CA)
      type="definitions" dict   → qualitative descriptions of letter grades
      type="regulation_section" → numbered clauses with optional keywords
      type="standards"          → grade distribution percentages
      type="methodology"        → evaluation formula and detail breakdown
    """
    results = []
    title   = item.get("title", "").strip()
    typ     = item.get("type",  "")
    content = item.get("content", {})

    if typ == "table":
        # Grading scale: score → letter → grade_point
        rows = []
        for row in content:
            if not isinstance(row, dict):
                continue
            score  = row.get("score", "")
            letter = row.get("letter_grade", "")
            gpa    = row.get("grade_point", "")
            rows.append(f"Score {score}: letter grade {letter}, GPA {gpa}.")
        if rows:
            text = f"{title}. " + " ".join(rows)
            results.append({"text": text, "source": src})

    elif typ == "definitions" and isinstance(content, list):
        # Special notations: W, WF, I, F, NR, CA, R
        for entry in content:
            if not isinstance(entry, dict):
                continue
            notation = entry.get("notation", "").strip()
            desc     = entry.get("description", "").strip()
            aliases  = entry.get("aliases", [])
            if not (notation and desc):
                continue
            aka  = ", ".join(aliases) if aliases else ""
            text = f"Grade notation {notation}"
            if aka:  text += f" (also known as: {aka})"
            text += f": {desc}"
            if len(text) >= MIN_CHUNK_LEN:
                results.append({"text": text, "source": src})

    elif typ == "definitions" and isinstance(content, dict):
        # Qualitative descriptions: A+_A, B+_B, etc.
        parts = [f"{title}."]
        for grade, desc in content.items():
            parts.append(f"{grade}: {desc}")
        text = " ".join(parts)
        if len(text) >= MIN_CHUNK_LEN:
            results.append({"text": text, "source": src})

    elif typ == "regulation_section":
        # Numbered clauses — each clause is its own chunk for precise retrieval
        for prov in content:
            if not isinstance(prov, dict):
                continue
            clause   = prov.get("clause", "").strip()
            body     = prov.get("content", prov.get("description", "")).strip()  # field is "content" in grading_en.json
            keywords = prov.get("keywords", [])
            if not body:
                continue
            text = f"{title}. Clause {clause}: {body}"
            if keywords:
                text += f" Related: {', '.join(keywords)}."
            if len(text) > CHUNK_SIZE:
                results.extend(split_chunks(text, src))
            elif len(text) >= MIN_CHUNK_LEN:
                results.append({"text": text, "source": src})

    elif typ == "standards":
        # Grade distribution percentages
        parts = [f"{title}."]
        for grade, pct in content.items():
            parts.append(f"{grade}: {pct}")
        text = " ".join(parts)
        if len(text) >= MIN_CHUNK_LEN:
            results.append({"text": text, "source": src})

    elif typ == "methodology":
        # Evaluation formula + breakdown
        formula = item.get("formula", "")
        details = item.get("details", {})
        parts   = [f"{title}."]
        if formula:
            parts.append(f"Formula: {formula}.")
        for v in details.values():
            parts.append(str(v))
        text = " ".join(parts)
        if len(text) >= MIN_CHUNK_LEN:
            results.append({"text": text, "source": src})

    return results


# ── PARSER: level.json ────────────────────────────────────────────────────────
def parse_level(item: dict, src: str) -> list[dict]:
    """
    level.json structure:
      { type:"regulation_table", title, content:[{ program, levels:[{name, credit_hours}], keywords }] }

    Generates one chunk per program with all levels listed together so Claude
    can see the full credit-hour range and answer "what level am I at?" correctly.
    The field is "credit_hours" (not "credits") — this was the original bug.
    """
    title   = item.get("title", "").strip()
    results = []

    for prog in item.get("content", []):
        if not isinstance(prog, dict):
            continue
        program  = prog.get("program", "").strip()    # e.g. "Bachelor"
        keywords = prog.get("keywords", [])
        levels   = prog.get("levels", [])

        level_parts = []
        for lvl in levels:
            if not isinstance(lvl, dict):
                continue
            name    = lvl.get("name", "").strip()          # e.g. "Level 3"
            credits = lvl.get("credit_hours", "").strip()  # e.g. "61-90" — correct field name
            if name and credits:
                level_parts.append(f"{name}: {credits} credit hours")

        if not level_parts:
            continue

        # All levels in one chunk — Claude needs the full table to compare ranges
        text = f"{title}. {program} program: " + "; ".join(level_parts) + "."
        if keywords:
            text += f" Related: {', '.join(keywords)}."

        if len(text) >= MIN_CHUNK_LEN:
            results.append({"text": text, "source": src})

    return results


# ── PARSER: schedule.json ─────────────────────────────────────────────────────
def parse_schedule(item: dict, src: str) -> list[dict]:
    """
    schedule.json structure:
      { instructor, course, index, semester, day, time, room, text, source }

    Each item already has a pre-built "text" field — use it directly.
    This is the richest representation and avoids reconstruction errors.
    """
    text = item.get("text", "").strip()
    if text and len(text) >= MIN_CHUNK_LEN:
        return [{"text": text, "source": src}]
    return []


# ── PARSER: teachers.json ─────────────────────────────────────────────────────
def parse_teacher(item: dict, src: str) -> list[dict]:
    """
    teachers.json structure:
      { name, department, room_number, phone, email }

    Skips fields whose value is "Unknown" — no point embedding unknown data.
    """
    def get(field: str) -> str:
        """Return the value only if it's non-empty and not a placeholder."""
        v = item.get(field, "")
        if v and str(v).strip().lower() not in ["unknown", "not available", "none", "null", ""]:
            return str(v).strip()
        return ""

    name = get("name")
    if not name:
        return []

    parts = [f"Teacher: {name}."]
    if dept  := get("department"):   parts.append(f"Department: {dept}.")
    if room  := get("room_number"):  parts.append(f"Room: {room}.")
    if phone := get("phone"):        parts.append(f"Phone: {phone}.")
    if email := get("email"):        parts.append(f"Email: {email}.")

    text = " ".join(parts)
    if len(text) >= MIN_CHUNK_LEN:
        return [{"text": text, "source": src}]
    return []


# ── PARSER: tuition.json ──────────────────────────────────────────────────────
def parse_tuition(item: dict, src: str) -> list[dict]:
    """
    tuition.json — two entry types:

    Type 1 — payment instructions:
      { type:"payment_instructions", school, bank, account_number,
        instructions (or text) }
      → Use the pre-built "instructions" or "text" field directly.

    Type 2 — tuition rates:
      { academic_year, enrollment_period (or admission_cohort), school,
        tuition_per_credit:{general_foundation, major_foundation} (or tuition:{...}), text }
      → Use the pre-built "text" field or reconstruct from fields.
    """
    # Try pre-built text fields first
    text = item.get("instructions", item.get("text", "")).strip()
    if text and len(text) >= MIN_CHUNK_LEN:
        return [{"text": text, "source": src}]

    # Fallback: reconstruct from fields if no pre-built text
    academic_year     = item.get("academic_year", "")
    enrollment_period = item.get("enrollment_period", item.get("admission_cohort", ""))
    school            = item.get("school", "")
    # Support both tuition_per_credit (new) and tuition (old)
    tuition_block     = item.get("tuition_per_credit", item.get("tuition", {}))
    general           = tuition_block.get("general_foundation", "")
    major             = tuition_block.get("major_foundation", "")

    if academic_year and school:
        parts = [f"Tuition fees for {academic_year} academic year."]
        if enrollment_period: parts.append(f"Enrollment period: {enrollment_period}.")
        if school:            parts.append(f"School: {school}.")
        if general:           parts.append(f"General foundation course: {int(general):,} MNT per credit.")
        if major:             parts.append(f"Major foundation course: {int(major):,} MNT per credit.")
        text = " ".join(parts)
        if len(text) >= MIN_CHUNK_LEN:
            return [{"text": text, "source": src}]

    return []


# ── MAIN JSON ROUTER ──────────────────────────────────────────────────────────
def read_json(filepath: str) -> list[dict]:
    """
    Route each JSON file to the correct parser based on its structure.
    Detects the file type from field names/values rather than the filename,
    so renaming files won't break ingestion.
    """
    src = os.path.basename(filepath)
    try:
        with open(filepath, encoding="utf-8") as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"  ❌ JSON error ({src}): {e}")
        return []

    items   = data if isinstance(data, list) else [data]
    results = []

    for item in items:
        if not isinstance(item, dict):
            continue

        # ── level.json: regulation_table ──────────────────
        if item.get("type") == "regulation_table":
            results.extend(parse_level(item, src))

        # ── grading.json: typed grading structures ─────────
        elif item.get("type") in ("table", "definitions", "regulation_section", "standards", "methodology"):
            chunks = parse_grading(item, src)
            for t in chunks:
                if len(t["text"]) > CHUNK_SIZE:
                    results.extend(split_chunks(t["text"], src))
                else:
                    results.append(t)

        # ── chuluu.json: provisions array ──────────────────
        elif "provisions" in item:
            results.extend(parse_chuluu(item, src))

        # ── tuition.json: payment instruction or rate ──────
        elif item.get("type") in ("payment_instruction", "payment_instructions") or "academic_year" in item:
            results.extend(parse_tuition(item, src))

        # ── schedule.json: pre-built text field ────────────
        elif "instructor" in item and "text" in item:
            results.extend(parse_schedule(item, src))

        # ── teachers.json: teacher contact info ────────────
        elif "name" in item and "department" in item and "email" in item:
            results.extend(parse_teacher(item, src))

        # ── courses.json: course info ──────────────────────
        elif "Course_Index" in item or "Course_Name" in item:
            results.extend(parse_course(item, src))

        # ── Fallback: unknown structure ─────────────────────
        else:
            text = "; ".join(f"{k}: {v}" for k, v in item.items() if v)
            if len(text) >= MIN_CHUNK_LEN:
                results.append({"text": text, "source": src})

    return results


# ── DOCX READER ───────────────────────────────────────────────────────────────
def read_docx(filepath: str) -> list[dict]:
    """Extract paragraphs and table cells from a Word document."""
    doc   = Document(filepath)
    src   = os.path.basename(filepath)
    lines = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            seen  = set()
            for c in row.cells:
                ct = c.text.strip()
                if ct and ct not in seen:
                    cells.append(ct)
                    seen.add(ct)
            if cells:
                lines.append(" | ".join(cells))

    return split_chunks("\n".join(lines), src)


# ── PDF READER ────────────────────────────────────────────────────────────────
def read_pdf(filepath: str) -> list[dict]:
    """Extract text from each PDF page and split into chunks."""
    src     = os.path.basename(filepath)
    results = []
    try:
        reader = PdfReader(filepath)
        for i, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            if text:
                results.extend(split_chunks(text, f"{src} (p.{i+1})"))
    except Exception as e:
        print(f"  ❌ PDF error ({src}): {e}")
    return results


# ── CSV READER ────────────────────────────────────────────────────────────────
def read_csv(filepath: str) -> list[dict]:
    """Convert CSV rows to text, batching 5 rows per chunk."""
    src     = os.path.basename(filepath)
    batch   = []
    results = []
    try:
        with open(filepath, encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                line = ", ".join(
                    f"{k}: {v.strip()}"
                    for k, v in row.items()
                    if v and v.strip()
                )
                if line:
                    batch.append(line)
                if len(batch) >= 5:
                    results.extend(split_chunks("\n".join(batch), src))
                    batch = []
        if batch:
            results.extend(split_chunks("\n".join(batch), src))
    except Exception as e:
        print(f"  ❌ CSV error ({src}): {e}")
    return results


# ── FILE DISPATCH ─────────────────────────────────────────────────────────────
READERS = {
    ".docx": read_docx,
    ".pdf":  read_pdf,
    ".json": read_json,
    ".csv":  read_csv,
}


def read_file(filepath: str) -> list[dict]:
    ext    = Path(filepath).suffix.lower()
    reader = READERS.get(ext)
    if reader is None:
        return []
    try:
        return reader(filepath)
    except Exception as e:
        print(f"  ❌ Error ({os.path.basename(filepath)}): {e}")
        return []


# ── PINECONE UPSERT ───────────────────────────────────────────────────────────
def upsert_all(chunks: list[dict], idx):
    """Encode all chunks and upload to Pinecone in batches."""
    total = len(chunks)
    print(f"🔄 Generating embeddings for {total} chunks...")

    texts      = [c["text"] for c in chunks]
    embeddings = embedder.encode(
        texts,
        batch_size=64,
        show_progress_bar=True,
        convert_to_numpy=True,
    )

    vectors = []
    for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
        vectors.append({
            "id":     make_id(chunk["text"], i),
            "values": emb.tolist(),
            "metadata": {
                "text":   chunk["text"][:1500],   # Pinecone metadata limit
                "source": chunk.get("source", ""),
            },
        })

    print(f"📤 Uploading to Pinecone ({len(vectors)} vectors)...")
    success = 0
    for i in range(0, len(vectors), BATCH_SIZE):
        batch = vectors[i:i + BATCH_SIZE]
        try:
            idx.upsert(vectors=batch)
            success += len(batch)
        except Exception as e:
            print(f"  ⚠️  Batch {i//BATCH_SIZE} error: {e}")
        time.sleep(0.15)

    print(f"✅ {success}/{len(vectors)} vectors uploaded successfully.")


# ── BM25 REBUILD ──────────────────────────────────────────────────────────────
def rebuild_bm25(chunks: list[dict]):
    """
    Rebuild and save the BM25 keyword index from the same chunks used for Pinecone.
    Must be called after ingestion so retrieval.py can load it on next app start.
    """
    try:
        from retrieval import bm25_index
        bm25_index.build(chunks)
        bm25_index.save()
        print(f"💾 BM25 index saved ({len(chunks)} docs).")
    except Exception as e:
        print(f"⚠️  BM25 rebuild failed: {e}")


# ── MAIN ─────────────────────────────────────────────────────────────────────
def start_ingestion():
    if not os.path.exists(DATA_FOLDER):
        os.makedirs(DATA_FOLDER)
        print(f"'{DATA_FOLDER}' folder created. Please copy your data files into it.")
        return

    idx = get_or_create_index()

    print("🗑️  Deleting existing vectors...")
    try:
        idx.delete(delete_all=True, namespace="")
        time.sleep(2)
    except Exception as e:
        print(f"  ⚠️  Error during deletion (may already be empty): {e}")

    all_chunks = []
    supported  = set(READERS.keys())

    print(f"\n📂 Reading files from '{DATA_FOLDER}'...")
    for filename in sorted(os.listdir(DATA_FOLDER)):
        ext = Path(filename).suffix.lower()
        if ext not in supported:
            continue
        filepath = os.path.join(DATA_FOLDER, filename)
        chunks   = read_file(filepath)
        if chunks:
            all_chunks.extend(chunks)
            print(f"  📄 {filename} → {len(chunks)} chunks")
        else:
            print(f"  ⚠️  {filename} → no chunks found")

    if not all_chunks:
        print("\n❌ No data to ingest.")
        return

    # Deduplicate by first 200 chars of text
    seen, unique = set(), []
    for c in all_chunks:
        key = c["text"][:200]
        if key not in seen:
            seen.add(key)
            unique.append(c)

    removed = len(all_chunks) - len(unique)
    print(f"\n📦 Total: {len(all_chunks)} | Duplicates removed: {removed} | Clean: {len(unique)}")

    upsert_all(unique, idx)
    rebuild_bm25(unique)

    print("\n🎯 Ingestion complete!")
    try:
        stats = idx.describe_index_stats()
        print(f"📊 Pinecone total vectors: {stats.total_vector_count}")
    except Exception:
        pass


if __name__ == "__main__":
    start_ingestion()