##60 кредит ямар түвшин бэ?
import re
import json
import os
from docx import Document

INPUT_FILE  = "журам.docx"
OUTPUT_FILE = "data/журам_chunks.json"

# ── Хэсэг тодорхойлох pattern ──────────────────────────
SECTION_RE = re.compile(
    r'^('
    r'\d+\.\d+(?:\.\d+)*\.?\s'   # 4.1.  4.3.7.  гэх мэт
    r'|НЭГ\.\s'
    r'|ХОЁР\.\s'
    r'|ГУРАВ\.\s'
    r'|ДӨРӨВ\.\s'
    r'|ТАВ\.\s'
    r'|ЗУРГАА\.\s'
    r'|ДОЛОО\.\s'
    r'|Хавсралт\s\d+'
    r')'
)

# Дээд түвшний гарчиг (гараар тодорхойлсон)
CHAPTER_MAP = {
    "НЭГ":    "1. Нийтлэг үндэслэл",
    "ХОЁР":   "2. Журмын нэр томьёо",
    "ГУРАВ":  "3. Сургалтын үйл ажиллагааны зохион байгуулалт",
    "ДӨРӨВ":  "4. Сургалтын үйл ажиллагаа",
    "ТАВ":    "5. Суралцагч чөлөө авах, үргэлжлүүлэн суралцах, сургуулиас чөлөөлөх, хасах",
    "ЗУРГАА": "6. Сургалтын төлбөр, хураамж",
    "ДОЛОО":  "7. Бусад",
}

def get_chapter(section_num: str) -> str:
    """Заалтын дугаараас бүлгийн нэрийг тодорхойлно."""
    if not section_num:
        return "Удиртгал"
    first = section_num.split(".")[0]
    for key, val in CHAPTER_MAP.items():
        if key[:3] == first[:3]:
            return val
    try:
        n = int(first)
        return f"Бүлэг {n}"
    except ValueError:
        return first


def extract_table_chunks(table, table_index: int) -> list[dict]:
    """Хүснэгтийг утга бүхий chunk болгоно."""
    rows = []
    headers = []

    for i, row in enumerate(table.rows):
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        # Давхардсан нүдийг хасна (merged cells)
        cells = list(dict.fromkeys(cells))
        cells = [c for c in cells if c]
        if not cells:
            continue
        if i == 0:
            headers = cells
        else:
            rows.append(cells)

    if not rows:
        return []

    chunks = []

    # Хүснэгт бүхэлдээ нэг chunk
    lines = []
    if headers:
        lines.append(" | ".join(headers))
        lines.append("-" * 40)
    for row in rows:
        lines.append(" | ".join(row))

    table_titles = {
        0: "Суралцагчийн түвшин тодорхойлох хүснэгт (нийт цуглуулсан багц цагаар)",
        1: "Хичээлийн үнэлгээний хүснэгт: хувьчилсан оноо → үсгэн дүн → тоон дүн",
        2: "Дүнгийн хуваарилалтын стандарт (хөндлөнгийн шалгалтын шалгуур)",
        3: "Лавлагаа, бичиг баримт олгох хугацааны хүснэгт",
        4: "Үсгэн тэмдэглэгээний тайлбар хүснэгт (NA, G, W, WF, NR, R, I, E, CA, CR, RC)",
    }
    title = table_titles.get(table_index, f"Хүснэгт {table_index + 1}")

    chunks.append({
        "id":      f"table-{table_index}",
        "section": title,
        "chapter": "Хүснэгт",
        "text":    f"{title}:\n" + "\n".join(lines),
    })

    # Мөр бүрийг тусдаа chunk (хайлтад илүү нарийн)
    if headers and len(rows) > 1:
        for j, row in enumerate(rows):
            if len(row) >= 2:
                row_text = f"{title} — {headers[0]}: {row[0]}"
                for k, h in enumerate(headers[1:], 1):
                    if k < len(row):
                        row_text += f", {h}: {row[k]}"
                chunks.append({
                    "id":      f"table-{table_index}-row-{j}",
                    "section": title,
                    "chapter": "Хүснэгт",
                    "text":    row_text,
                })

    return chunks


def parse_журам(filepath: str) -> list[dict]:
    doc = Document(filepath)
    chunks: list[dict] = []
    table_map: dict[int, list] = {}  # paragraph index → table chunks

    # Хүснэгтүүдийг урьдчилан боловсруулна
    for i, table in enumerate(doc.tables):
        table_map[i] = extract_table_chunks(table, i)

    # Paragraph-уудыг хэсэглэнэ
    current_num   = ""
    current_title = ""
    current_body  = []
    table_idx     = 0

    def flush(num, title, body_lines):
        body = " ".join(body_lines).strip()
        if len(body) < 20:
            return
        chapter = get_chapter(num)
        label   = f"[{num}] {title}" if num else title
        chunks.append({
            "id":      f"section-{num or 'intro'}",
            "section": label,
            "chapter": chapter,
            "text":    f"{label}\n{body}",
        })

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "tbl":
            # Хүснэгт олдвол flush хийж, хүснэгтийн chunk нэмнэ
            flush(current_num, current_title, current_body)
            current_body = []
            if table_idx < len(doc.tables):
                tbl_chunks = extract_table_chunks(doc.tables[table_idx], table_idx)
                chunks.extend(tbl_chunks)
                table_idx += 1
            continue

        if tag != "p":
            continue

        # Paragraph-ын текстийг авна
        text = "".join(
            node.text or ""
            for node in element.iter()
            if node.tag.endswith("}t")
        ).strip()

        if not text:
            continue

        match = SECTION_RE.match(text)
        if match:
            flush(current_num, current_title, current_body)
            current_num   = match.group(1).strip().rstrip(".")
            current_title = text[:120]
            current_body  = [text]
        else:
            current_body.append(text)

    flush(current_num, current_title, current_body)

    # ── Нэмэлт: ойролцоо заалтуудыг нэгтгэсэн chunk ──
    # (жишээ нь: 4.1.1–4.1.5 нэг бүлэг болгоно)
    merged = []
    window = []
    prev_chapter = ""
    for c in chunks:
        if c["chapter"] == "Хүснэгт":
            if window:
                merged.append(merge_window(window))
                window = []
            merged.append(c)
            prev_chapter = c["chapter"]
            continue

        if c["chapter"] != prev_chapter and window:
            merged.append(merge_window(window))
            window = []

        window.append(c)
        prev_chapter = c["chapter"]

        if len(window) >= 5:
            merged.append(merge_window(window))
            window = []

    if window:
        merged.append(merge_window(window))

    return merged


def merge_window(window: list[dict]) -> dict:
    """5 ойролцоо chunk-г нэг том chunk болгоно (бүтцийг алдахгүй)."""
    if len(window) == 1:
        return window[0]
    combined = "\n\n".join(c["text"] for c in window)
    ids      = " — ".join(c["id"] for c in window)
    return {
        "id":      ids,
        "section": f"{window[0]['section']} … {window[-1]['section']}",
        "chapter": window[0]["chapter"],
        "text":    combined,
    }


def main():
    if not os.path.exists(INPUT_FILE):
        print(f"❌ '{INPUT_FILE}' файл олдсонгүй. data/ хавтастай ижил газар тавина уу.")
        return

    os.makedirs("data", exist_ok=True)

    print(f"📄 '{INPUT_FILE}' уншиж байна...")
    chunks = parse_журам(INPUT_FILE)

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)

    print(f"✅ {len(chunks)} chunk → '{OUTPUT_FILE}'")

    # Дээжүүд харуулна
    print("\n── Дээж chunk-ууд ──")
    samples = [c for c in chunks if "4.1" in c["id"] or "table" in c["id"]][:4]
    for s in samples:
        print(f"\n[{s['id']}] {s['section'][:60]}")
        print(f"  {s['text'][:200]}...")


if __name__ == "__main__":
    main()